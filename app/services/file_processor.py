from fastapi import UploadFile
import PyPDF2
import docx
import io
import logging

# Set up logging
logger = logging.getLogger(__name__)

async def process_pdf(file: UploadFile) -> str:
    try:
        logger.info(f"Processing PDF file: {file.filename}")
        # Read the uploaded file into memory
        contents = await file.read()
        pdf_file = io.BytesIO(contents)
        
        # Create PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        logger.info(f"PDF loaded with {len(pdf_reader.pages)} pages")
        
        # Extract text from all pages
        text = ""
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
            logger.debug(f"Extracted {len(page_text)} characters from page {i+1}")
        
        logger.info(f"PDF processing complete. Total extracted text: {len(text)} characters")
        return text.strip()
    except Exception as e:
        logger.error(f"Error processing PDF {file.filename}: {str(e)}", exc_info=True)
        raise Exception(f"Error processing PDF: {str(e)}")

async def process_docx(file: UploadFile) -> str:
    try:
        logger.info(f"Processing DOCX file: {file.filename}")
        # Read the uploaded file into memory
        contents = await file.read()
        doc_file = io.BytesIO(contents)
        
        # Create document object
        doc = docx.Document(doc_file)
        logger.info(f"DOCX loaded with {len(doc.paragraphs)} paragraphs")
        
        # Extract text from all paragraphs
        text = ""
        for i, para in enumerate(doc.paragraphs):
            text += para.text + "\n"
            logger.debug(f"Extracted paragraph {i+1}: {len(para.text)} characters")
        
        logger.info(f"DOCX processing complete. Total extracted text: {len(text)} characters")
        return text.strip()
    except Exception as e:
        logger.error(f"Error processing DOCX {file.filename}: {str(e)}", exc_info=True)
        raise Exception(f"Error processing DOCX: {str(e)}")

async def process_txt(file: UploadFile) -> str:
    try:
        logger.info(f"Processing TXT file: {file.filename}")
        # Read the uploaded file into memory
        contents = await file.read()
        text = contents.decode('utf-8').strip()
        logger.info(f"TXT processing complete. Total extracted text: {len(text)} characters")
        return text
    except Exception as e:
        logger.error(f"Error processing TXT {file.filename}: {str(e)}", exc_info=True)
        raise Exception(f"Error processing TXT: {str(e)}")

async def process_file(file: UploadFile) -> str:
    try:
        logger.info(f"Starting file processing for: {file.filename}")
        if file.filename.endswith('.pdf'):
            text = await process_pdf(file)
        elif file.filename.endswith('.docx'):
            text = await process_docx(file)
        elif file.filename.endswith('.txt'):
            text = await process_txt(file)
        else:
            error_msg = f"Unsupported file type: {file.filename}"
            logger.error(error_msg)
            raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT files.")
        
        logger.info(f"File processing complete for {file.filename}. Text length: {len(text)} characters")
        return text
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {str(e)}", exc_info=True)
        raise Exception(f"Error processing file: {str(e)}") 